# -*- coding: utf-8 -*-

import re
import time
import socket
import os
import shutil
import sys
from urllib.parse import urlparse
from dataclasses import dataclass
from io import BytesIO
from typing import List, Optional

import pandas as pd
import streamlit as st
from docx import Document

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException


URL = "https://chat.sonax.net.br/app/omnichannel/chat"
PAGE_LOAD_TIMEOUT_S = 45.0
HEADLESS_LOGIN_TIMEOUT_S = 45.0


@dataclass
class Cliente:
    nome: str
    placa: str
    telefone: str
    endereco: str = ""
    horario: str = ""  # data


PHONE_RE = re.compile(r"(?:\+?55)?\s*\(?\d{2}\)?\s*\d{4,5}-?\d{4}")
PLATE_CANDIDATE_RE = re.compile(r"\b[A-Z]{3}[A-Z0-9]{4,5}\b", re.IGNORECASE)


def normalize_phone(raw: str) -> Optional[str]:
    if not raw:
        return None
    digits = re.sub(r"\D+", "", raw)
    if len(digits) == 11:
        return digits
    if len(digits) == 13 and digits.startswith("55"):
        return digits
    return None


def phone_variations(phone: str) -> List[str]:
    d = re.sub(r"\D+", "", phone)
    out = []
    if len(d) == 11:
        out += [d, "55" + d]
    elif len(d) == 13 and d.startswith("55"):
        out += [d, d[2:]]
    else:
        out += [d]
    seen, res = set(), []
    for x in out:
        if x and x not in seen:
            seen.add(x)
            res.append(x)
    return res


def normalize_plate(raw: str) -> str:
    if not raw:
        return ""
    return re.sub(r"[^A-Za-z0-9]+", "", raw).upper()


def is_valid_plate_relaxed(s: str) -> bool:
    s = normalize_plate(s)
    if not (7 <= len(s) <= 8):
        return False
    if not re.match(r"^[A-Z]{3}[A-Z0-9]+$", s):
        return False
    if not any(ch.isdigit() for ch in s):
        return False
    return True


def find_plate_in_text(text: str) -> str:
    m = re.search(r"(?i)\bplaca\s*:\s*([A-Z0-9-]+)", text)
    if m:
        p = normalize_plate(m.group(1))
        if is_valid_plate_relaxed(p):
            return p
    for cand in PLATE_CANDIDATE_RE.findall(text):
        p = normalize_plate(cand)
        if is_valid_plate_relaxed(p):
            return p
    return ""


def docx_lines_preserve_blanks(doc: Document) -> List[str]:
    out: List[str] = []
    for p in doc.paragraphs:
        out.append((p.text or "").strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            out.append("" if not any(cells) else " | ".join(cells))
    return out


def parse_record_block(block_lines: List[str]) -> Optional[Cliente]:
    lines = [ln.strip() for ln in block_lines if ln and ln.strip()]
    if not lines:
        return None

    joined = "\n".join(lines)

    mph = PHONE_RE.search(joined)
    phone = normalize_phone(mph.group(0)) if mph else None
    plate = find_plate_in_text(joined)

    # nome
    nome = ""
    for ln in lines[:6]:
        if phone and phone[-11:] in re.sub(r"\D+", "", ln):
            continue
        if "placa" in ln.lower():
            continue
        if re.search(r"(?i)\b(rua|avenida|av\.|rodovia|estrada)\b", ln):
            continue
        nome = ln
        break
    if not nome:
        nome = "Sem nome"

    # endereço (primeira linha com rua/av)
    endereco = ""
    for ln in lines:
        if re.search(r"(?i)\b(rua|avenida|av\.|rodovia|estrada)\b", ln):
            endereco = ln
            break

    # data (dd/mm/aaaa)
    data = ""
    mdate = re.search(r"\b\d{2}/\d{2}/\d{4}\b", joined)
    if mdate:
        data = mdate.group(0)

    if phone and plate:
        return Cliente(nome=nome, telefone=phone, placa=plate, endereco=endereco, horario=data)

    return None


def parse_docx_to_clients(file_bytes: bytes) -> List[Cliente]:
    doc = Document(BytesIO(file_bytes))
    lines = docx_lines_preserve_blanks(doc)

    blocks, cur = [], []
    for ln in lines:
        if ln.strip() == "":
            if cur:
                blocks.append(cur)
                cur = []
        else:
            cur.append(ln)
    if cur:
        blocks.append(cur)

    out = []
    seen = set()
    for b in blocks:
        c = parse_record_block(b)
        if c:
            k = (c.telefone, c.placa)
            if k not in seen:
                seen.add(k)
                out.append(c)
    return out


def port_open(host: str, port: int, timeout_s: float = 0.35) -> bool:
    try:
        with socket.create_connection((host, port), timeout=timeout_s):
            return True
    except Exception:
        return False


def _is_linux() -> bool:
    return sys.platform.startswith("linux")


def _is_headless_server_runtime() -> bool:
    return _is_linux() and not os.getenv("DISPLAY")


def _supports_local_debug_attach() -> bool:
    if sys.platform.startswith(("win", "darwin")):
        return True
    if _is_linux() and os.getenv("DISPLAY"):
        return True
    return False


def _find_chromedriver_path() -> Optional[str]:
    env_path = (os.getenv("CHROMEDRIVER_PATH") or "").strip()
    if env_path:
        return env_path

    for p in (
        "/usr/bin/chromedriver",
        "/usr/lib/chromium/chromedriver",
        "/usr/lib/chromium-browser/chromedriver",
        "/snap/bin/chromedriver",
    ):
        if os.path.exists(p):
            return p

    system_path = shutil.which("chromedriver")
    if system_path and "/.cache/selenium/" not in system_path:
        return system_path
    return None


def _configure_linux_runtime(opts: Options) -> None:
    if not _is_linux():
        return
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--disable-software-rasterizer")
    opts.add_argument("--window-size=1920,1080")
    if _is_headless_server_runtime() or not os.getenv("DISPLAY"):
        opts.add_argument("--headless=new")

    chrome_bin = (
        (os.getenv("CHROME_BINARY") or "").strip()
        or (os.getenv("CHROMIUM_BINARY") or "").strip()
    )
    if chrome_bin:
        opts.binary_location = chrome_bin
        return

    for p in (
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/usr/bin/chromium-browser",
        "/usr/bin/chromium",
    ):
        if os.path.exists(p):
            opts.binary_location = p
            break


def _build_chrome_service() -> Optional[Service]:
    if not _is_linux():
        return None
    path = _find_chromedriver_path()
    return Service(executable_path=path) if path else None


def _start_chrome(opts: Options, service: Optional[Service] = None) -> webdriver.Chrome:
    if service:
        try:
            return webdriver.Chrome(service=service, options=opts)
        except WebDriverException:
            pass
    return webdriver.Chrome(options=opts)


# bloqueando popups
def make_driver_attach(debug_port: int) -> webdriver.Chrome:
    opts = Options()
    opts.add_experimental_option("debuggerAddress", f"127.0.0.1:{debug_port}")
    opts.add_argument("--disable-notifications")
    opts.add_argument("--disable-popup-blocking")
    _configure_linux_runtime(opts)
    return _start_chrome(opts, service=_build_chrome_service())


def make_driver_new() -> webdriver.Chrome:
    opts = Options()
    opts.add_argument("--start-maximized")
    opts.add_argument("--disable-notifications")
    opts.add_argument("--disable-popup-blocking")
    _configure_linux_runtime(opts)
    return _start_chrome(opts, service=_build_chrome_service())


def _read_secret_value(key: str) -> str:
    try:
        val = st.secrets.get(key, "")
        return str(val).strip() if val is not None else ""
    except Exception:
        return ""


def _read_nested_secret_value(section: str, key: str) -> str:
    try:
        block = st.secrets.get(section, {})
        if isinstance(block, dict):
            val = block.get(key, "")
            return str(val).strip() if val is not None else ""
    except Exception:
        pass
    return ""


def _get_headless_login_credentials() -> tuple[str, str]:
    user = (
        _read_secret_value("SONAX_USERNAME")
        or _read_secret_value("SONAX_USER")
        or _read_secret_value("SONAX_LOGIN")
        or _read_nested_secret_value("sonax", "username")
        or _read_nested_secret_value("sonax", "user")
        or _read_nested_secret_value("sonax", "login")
        or (os.getenv("SONAX_USERNAME") or "").strip()
        or (os.getenv("SONAX_USER") or "").strip()
        or (os.getenv("SONAX_LOGIN") or "").strip()
    )
    pwd = (
        _read_secret_value("SONAX_PASSWORD")
        or _read_secret_value("SONAX_PASS")
        or _read_nested_secret_value("sonax", "password")
        or _read_nested_secret_value("sonax", "pass")
        or (os.getenv("SONAX_PASSWORD") or "").strip()
        or (os.getenv("SONAX_PASS") or "").strip()
    )
    return user, pwd


def _headless_credentials_validation() -> tuple[str, str, List[str], List[str]]:
    user, pwd = _get_headless_login_credentials()
    missing = []
    hints = []
    if not user:
        missing.append("usuario")
        hints.append("SONAX_USERNAME (ou SONAX_USER / SONAX_LOGIN / [sonax].username)")
    if not pwd:
        missing.append("senha")
        hints.append("SONAX_PASSWORD (ou SONAX_PASS / [sonax].password)")
    return user, pwd, missing, hints


def _safe_get(driver, url: str, timeout_s: float = PAGE_LOAD_TIMEOUT_S) -> None:
    try:
        driver.set_page_load_timeout(timeout_s)
    except Exception:
        pass
    try:
        driver.get(url)
    except TimeoutException:
        pass


def _find_visible_input(driver, xpath: str):
    for el in driver.find_elements(By.XPATH, xpath):
        try:
            if el.is_displayed() and el.is_enabled():
                return el
        except Exception:
            continue
    return None


def _set_input_value(el, value: str):
    el.click()
    el.send_keys(Keys.CONTROL, "a")
    el.send_keys(Keys.BACKSPACE)
    el.send_keys(value or "")


def has_authenticated_sonax_session(driver, timeout_s: float = 6.0) -> bool:
    end_at = time.time() + timeout_s
    while time.time() < end_at:
        try:
            current_url = (driver.current_url or "").lower()
        except Exception:
            current_url = ""
        if "/login" not in current_url:
            return True
        time.sleep(0.2)
    return False


def try_headless_login_with_credentials(driver, timeout_s: float = HEADLESS_LOGIN_TIMEOUT_S) -> bool:
    user, pwd = _get_headless_login_credentials()
    if not user or not pwd:
        return False

    end_at = time.time() + timeout_s
    while time.time() < end_at:
        if has_authenticated_sonax_session(driver, timeout_s=0.8):
            return True

        user_input = _find_visible_input(
            driver,
            "//input[not(@type='password') and not(@type='hidden') and not(@disabled)]",
        )
        pwd_input = _find_visible_input(
            driver,
            "//input[@type='password' and not(@disabled)]",
        )
        if user_input is None or pwd_input is None:
            time.sleep(0.4)
            continue

        _set_input_value(user_input, user)
        _set_input_value(pwd_input, pwd)

        submit = _find_visible_input(
            driver,
            "//button[@type='submit' or contains(normalize-space(.),'Entrar') or contains(normalize-space(.),'Login')]",
        )
        if submit is not None:
            submit.click()
        else:
            pwd_input.send_keys(Keys.ENTER)

        if has_authenticated_sonax_session(driver, timeout_s=8.0):
            return True
        time.sleep(0.4)
    return False


def _url_host(raw_url: str) -> str:
    try:
        return (urlparse(raw_url).hostname or "").strip()
    except Exception:
        return ""


def _host_reachable(host: str, timeout_s: float = 2.0) -> bool:
    if not host:
        return False
    for port in (443, 80):
        if port_open(host, port, timeout_s=timeout_s):
            return True
    return False


def maybe_close_popup(driver) -> bool:
    try:
        btn = WebDriverWait(driver, 1.2).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "#botao-fechar"))
        )
        btn.click()
        time.sleep(0.12)
        return True
    except Exception:
        return False


def click_retry(driver, by, value, tries=3, timeout=25):
    last = None
    for _ in range(tries):
        try:
            el = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by, value)))
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            el.click()
            return el
        except Exception as e:
            last = e
            maybe_close_popup(driver)
            time.sleep(0.25)
    raise last


def type_retry(driver, by, value, text, clear=True, press_enter=False, tries=3, timeout=25):
    last = None
    for _ in range(tries):
        try:
            el = WebDriverWait(driver, timeout).until(EC.visibility_of_element_located((by, value)))
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
            el.click()
            if clear:
                el.send_keys(Keys.CONTROL, "a")
                el.send_keys(Keys.BACKSPACE)
            el.send_keys(text)
            if press_enter:
                el.send_keys(Keys.ENTER)
            return el
        except Exception as e:
            last = e
            maybe_close_popup(driver)
            time.sleep(0.25)
    raise last


def ensure_sonax_tab(driver):
    target = "chat.sonax.net.br"
    for h in driver.window_handles:
        driver.switch_to.window(h)
        try:
            if target in (driver.current_url or ""):
                return
        except Exception:
            pass

    try:
        _safe_get(driver, URL)
    except Exception:
        driver.execute_script(f"window.open('{URL}','_blank');")
        driver.switch_to.window(driver.window_handles[-1])

    WebDriverWait(driver, 30).until(
        lambda d: d.execute_script("return document.readyState") in ("interactive", "complete")
    )


# Seleção
SEL_CONTATOS = (By.XPATH, "//a[contains(@class,'nav-link')][contains(.,'Contatos')]")
SEL_BUSCA = (By.CSS_SELECTOR, "input.form-control.input-search")
SEL_CONVERSAR = (By.CSS_SELECTOR, "button#dropdownBasic1")
SEL_LWSIMAPP = (By.XPATH, "//*[contains(@class,'ml-2') and contains(.,'LWSIMAPP')]")
SEL_COMBOBOX = (By.CSS_SELECTOR, "input[role='combobox']")
SEL_TEMPLATE = (
    By.XPATH,
    "//span[contains(@class,'ng-option-label') and normalize-space(.)='abertura_de_diagnostico_tagpro']",
)

# ⚠️ esse placeholder depende do texto EXATO do site. Com UTF-8 corrigido fica:
SEL_VAR_INPUTS = (By.CSS_SELECTOR, "input.form-control[placeholder='Insira a variável aqui']")
SEL_ENVIAR = (By.XPATH, "//button[contains(@class,'btn-primary') and normalize-space(.)='Enviar']")


def click_card_contact(driver, phone_digits: str) -> bool:
    digits = re.sub(r"\D+", "", phone_digits)
    xpath = (
        "//*[contains(@class,'kt-widget__info')]"
        f"[.//span[contains(@class,'kt-widget__desc')][contains(translate(., ' ()+-', ''), '{digits}')]]"
    )
    try:
        click_retry(driver, By.XPATH, xpath, tries=2, timeout=8)
        return True
    except Exception:
        try:
            click_retry(driver, By.CSS_SELECTOR, ".kt-widget__info", tries=1, timeout=3)
            return True
        except Exception:
            return False


def fill_template_variables_in_order(driver, placa: str, data: str, endereco: str):
    """
    Como os 3 campos são iguais, preenche pela ORDEM:
    [0]=placa, [1]=data, [2]=endereço
    """
    inputs = WebDriverWait(driver, 25).until(lambda d: d.find_elements(*SEL_VAR_INPUTS))
    if len(inputs) < 3:
        raise RuntimeError(f"Esperava 3 campos de variável, mas encontrei {len(inputs)}.")

    values = [placa or "", data or "", endereco or ""]
    for i in range(3):
        el = inputs[i]
        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
        el.click()
        el.send_keys(Keys.CONTROL, "a")
        el.send_keys(Keys.BACKSPACE)
        el.send_keys(values[i])
        time.sleep(0.1)


def run_one_client(driver, client: Cliente, log):
    maybe_close_popup(driver)

    log(f"➡️ {client.nome}: Contatos")
    click_retry(driver, *SEL_CONTATOS, tries=3, timeout=30)
    maybe_close_popup(driver)

    found = False
    for ph in phone_variations(client.telefone):
        log(f"🔎 {client.nome}: Buscar {ph}")
        click_retry(driver, *SEL_BUSCA, tries=3, timeout=30)
        type_retry(driver, *SEL_BUSCA, ph, clear=True, press_enter=True, tries=3, timeout=30)
        time.sleep(0.9)
        if click_card_contact(driver, ph):
            found = True
            break

    if not found:
        log(f"❌ {client.nome}: NÃO ENCONTRADO")
        return {"nome": client.nome, "telefone": client.telefone, "placa": client.placa, "status": "NÃO ENCONTRADO"}

    maybe_close_popup(driver)

    log(f"💬 {client.nome}: Conversar")
    click_retry(driver, *SEL_CONVERSAR, tries=3, timeout=30)
    maybe_close_popup(driver)

    log(f"📲 {client.nome}: LWSIMAPP")
    click_retry(driver, *SEL_LWSIMAPP, tries=3, timeout=30)
    maybe_close_popup(driver)

    log(f"🧾 {client.nome}: Template")
    click_retry(driver, *SEL_COMBOBOX, tries=3, timeout=30)
    click_retry(driver, *SEL_TEMPLATE, tries=3, timeout=30)
    maybe_close_popup(driver)

    log(f"⌨️ {client.nome}: preenchendo variáveis (placa/data/endereço)")
    fill_template_variables_in_order(driver, client.placa, client.horario, client.endereco)
    maybe_close_popup(driver)

    log(f"📨 {client.nome}: Enviar")
    click_retry(driver, *SEL_ENVIAR, tries=3, timeout=30)

    log(f"✅ {client.nome}: OK")
    return {"nome": client.nome, "telefone": client.telefone, "placa": client.placa, "status": "OK"}


# UI
st.set_page_config(page_title="Sonax Automação Kezia", layout="wide")
st.markdown(
    """
<style>
section[data-testid="stSidebar"] { display: none !important; }
div[data-testid="collapsedControl"] { display: none !important; }
.block-container { padding-top: 1.2rem; }

/* Fonte com suporte a emojis (principalmente Windows) */
html, body, [class*="st-"] {
  font-family: "Segoe UI", "Segoe UI Emoji", "Apple Color Emoji", "Noto Color Emoji", sans-serif;
}
</style>
""",
    unsafe_allow_html=True,
)

st.title("Sonax • Automação da KESIA")

st.session_state.setdefault("attach", not _is_headless_server_runtime())
st.session_state.setdefault("debug_port", 9222)
st.session_state.setdefault("max_items", 50)
runtime_mode = "deploy_headless" if _is_headless_server_runtime() else "local_interativo"

st.info(f"Modo de execução ativo: `{runtime_mode}`")

if runtime_mode == "deploy_headless":
    st.warning(
        "Ambiente de deploy detectado (sem interface gráfica). "
        "O Chrome roda em modo headless no servidor."
    )
    user, pwd, missing, hints = _headless_credentials_validation()
    if missing:
        st.error(
            "Credenciais do Sonax ausentes para o deploy. "
            f"Faltando: {', '.join(missing)}."
        )
        st.code("\n".join(hints))
    else:
        st.success("Credenciais do Sonax para deploy: configuradas.")

with st.expander("Configurar", expanded=False):
    st.session_state.max_items = st.number_input(
        "Processar quantos clientes?",
        1,
        500,
        int(st.session_state.max_items),
        1,
    )

st.markdown("---")

uploaded = st.file_uploader("📄 Envie o Arquivo DOCX com os clientes", type=["docx"])
if not uploaded:
    st.info("Envie o arquivo para carregar os clientes.")
    st.stop()

clients = parse_docx_to_clients(uploaded.getvalue())
clients = clients[: int(st.session_state.max_items)]

st.success(f"Clientes carregados: {len(clients)}")

with st.expander("Ver clientes identificados", expanded=False):
    st.dataframe(pd.DataFrame([c.__dict__ for c in clients]), use_container_width=True, hide_index=True)

start = st.button("▶ Iniciar automação", type="primary")

if start:
    status_box = st.empty()
    log_box = st.empty()
    prog = st.progress(0)
    logs = []

    def log(msg: str):
        logs.append(msg)
        log_box.write("\n".join(logs[-35:]))

    results = []
    driver = None

    try:
        status_box.info("Validando pré-requisitos...")
        if not _supports_local_debug_attach() and st.session_state.attach:
            status_box.info("Deploy sem interface gráfica: desativando attach local.")
            st.session_state.attach = False

        if runtime_mode == "deploy_headless":
            user, pwd, missing, hints = _headless_credentials_validation()
            if missing:
                raise RuntimeError(
                    "Credenciais do Sonax ausentes no deploy. "
                    f"Faltando: {', '.join(missing)}. "
                    f"Chaves aceitas: {'; '.join(hints)}."
                )

        if st.session_state.attach:
            status_box.info("Testando porta do Chrome...")
            if not port_open("127.0.0.1", int(st.session_state.debug_port)):
                status_box.warning("Não tem Chrome na porta informada. Vou abrir um Chrome novo.")
                driver = make_driver_new()
                _safe_get(driver, URL)
            else:
                status_box.info("Conectando ao Chrome existente...")
                driver = make_driver_attach(int(st.session_state.debug_port))
        else:
            if runtime_mode == "deploy_headless":
                status_box.info("Iniciando Chromium headless no servidor...")
                host = _url_host(URL)
                if not _host_reachable(host):
                    raise RuntimeError(
                        f"O servidor de deploy não alcança {host}. "
                        "Se depende de rede interna/VPN, execute localmente."
                    )
                if not _find_chromedriver_path():
                    raise RuntimeError(
                        "ChromeDriver não encontrado no deploy. "
                        "Confirme packages.txt com chromium + chromium-driver."
                    )
            else:
                status_box.info("Abrindo Chrome novo...")
            driver = make_driver_new()
            _safe_get(driver, URL)

        status_box.info("Indo para a aba do Sonax...")
        ensure_sonax_tab(driver)

        if runtime_mode == "deploy_headless" and not has_authenticated_sonax_session(driver):
            status_box.info("Sessão não autenticada no deploy. Tentando login automático...")
            if not try_headless_login_with_credentials(driver):
                raise RuntimeError(
                    "Não foi possível autenticar no Sonax em modo deploy. "
                    "Valide SONAX_USERNAME/SONAX_PASSWORD em Settings > Secrets."
                )

        status_box.success("Executando automação...")
        for i, c in enumerate(clients, start=1):
            status_box.info(f"Processando {i}/{len(clients)}: {c.nome}")
            r = run_one_client(driver, c, log)
            results.append(r)
            prog.progress(i / len(clients))

        status_box.success("Finalizado!")
    except Exception as e:
        status_box.error(f"Erro na automação: {e}")
        log("⚠️ Se o template não tiver 3 variáveis ou a ordem for diferente, me avise que eu mapeio pelo label.")
    finally:
        pass

    if results:
        rdf = pd.DataFrame(results)
        st.subheader("Resultado")
        st.dataframe(rdf, use_container_width=True, hide_index=True)
        st.download_button(
            "Baixar resultado (.csv)",
            data=rdf.to_csv(index=False).encode("utf-8-sig"),  # BOM pro Excel
            file_name="resultado_sonax.csv",
            mime="text/csv",
        )
